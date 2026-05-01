import io
import os
import tempfile
from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import Dict, Any, List

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Inches
from fpdf import FPDF

router = APIRouter()

class ExportRequest(BaseModel):
    data: Dict[str, Any]

def generate_charts(analysis: Dict[str, Any], temp_dir: str) -> Dict[str, str]:
    images = {}
    
    # Sentiment
    sentiments = analysis.get("sentiments", {})
    if sentiments:
        plt.figure(figsize=(5, 3.5))
        plt.pie(list(sentiments.values()), labels=list(sentiments.keys()), autopct='%1.1f%%', startangle=140)
        plt.title('Sentiment Analysis')
        sentiment_path = os.path.join(temp_dir, 'sentiment.png')
        plt.savefig(sentiment_path, bbox_inches='tight')
        plt.close()
        images['sentiment'] = sentiment_path
        
    # Organizations
    organizations = analysis.get("organizations", {})
    if organizations:
        plt.figure(figsize=(5, 3.5))
        plt.bar(list(organizations.keys()), list(organizations.values()), color='#3b82f6')
        plt.title('Top Organizations')
        plt.ylabel('Count')
        plt.xticks(rotation=45, ha="right")
        org_path = os.path.join(temp_dir, 'org.png')
        plt.savefig(org_path, bbox_inches='tight')
        plt.close()
        images['org'] = org_path
        
    # Persons
    persons = analysis.get("persons", {})
    if persons:
        plt.figure(figsize=(5, 3.5))
        plt.bar(list(persons.keys()), list(persons.values()), color='#10b981')
        plt.title('Top People Mentioned')
        plt.ylabel('Count')
        plt.xticks(rotation=45, ha="right")
        person_path = os.path.join(temp_dir, 'person.png')
        plt.savefig(person_path, bbox_inches='tight')
        plt.close()
        images['person'] = person_path
        
    # Topics
    topics = analysis.get("topics", [])
    if topics:
        labels = [t.get("topic") for t in topics]
        counts = [t.get("count") for t in topics]
        plt.figure(figsize=(6, 3.5))
        plt.bar(labels, counts, color='#8b5cf6')
        plt.title('Top Topics')
        plt.ylabel('Count')
        plt.xticks(rotation=45, ha="right")
        topics_path = os.path.join(temp_dir, 'topics.png')
        plt.savefig(topics_path, bbox_inches='tight')
        plt.close()
        images['topics'] = topics_path

    return images

@router.post("/export/word")
def export_word(req: ExportRequest):
    data = req.data
    summary = data.get("summary", {})
    analysis = data.get("analysis", {})
    
    with tempfile.TemporaryDirectory() as tmpdir:
        chart_paths = generate_charts(analysis, tmpdir)
        
        doc = Document()
        doc.add_heading('YouTube Data Analysis Report', 0)
        
        doc.add_heading('Summary', level=1)
        doc.add_paragraph(f"Total Videos Analyzed: {summary.get('total_videos', 0)}")
        doc.add_paragraph(f"Total Comments Processed: {summary.get('total_comments_analyzed', 0)}")
        
        doc.add_heading('Methodology', level=1)
        doc.add_paragraph("Data extracted from YouTube Data API dynamically based on user search query. Comments processed via NLP, cleaned, and analyzed for named entities using Spacy. Overall sentiment analyzed via VADER.")
        
        doc.add_heading('Top topics/keywords', level=1)
        topics = analysis.get("topics", [])
        for t in topics:
            doc.add_paragraph(f"{t.get('topic')}: {t.get('count')}")
        if 'topics' in chart_paths:
            doc.add_picture(chart_paths['topics'], width=Inches(4))

        doc.add_heading('Named entities frequency', level=1)
        organizations = analysis.get("organizations", {})
        doc.add_paragraph("Top Organizations:")
        for k, v in organizations.items():
            doc.add_paragraph(f"{k}: {v}")
        if 'org' in chart_paths:
            doc.add_picture(chart_paths['org'], width=Inches(4))
            
        persons = analysis.get("persons", {})
        doc.add_paragraph("Top Persons:")
        for k, v in persons.items():
            doc.add_paragraph(f"{k}: {v}")
        if 'person' in chart_paths:
            doc.add_picture(chart_paths['person'], width=Inches(4))
            
        doc.add_heading('Sentiment charts', level=1)
        sentiments = analysis.get("sentiments", {})
        for k, v in sentiments.items():
            doc.add_paragraph(f"{str(k).capitalize()}: {v}")
        if 'sentiment' in chart_paths:
            doc.add_picture(chart_paths['sentiment'], width=Inches(4))
            
        doc.add_heading('Key insights', level=1)
        doc.add_paragraph("1. Discourse is varied depending on dominant entities.")
        doc.add_paragraph("2. Sentiment splits highlight community alignment.")
        
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
    
    return StreamingResponse(
        file_stream, 
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=analysis_report.docx"}
    )

class PDF(FPDF):
    def header(self):
        self.set_font('helvetica', 'B', 15)
        self.cell(0, 10, 'YouTube Data Analysis Report', ln=True, align='C')
        self.ln(5)

@router.post("/export/pdf")
def export_pdf(req: ExportRequest):
    data = req.data
    summary = data.get("summary", {})
    analysis = data.get("analysis", {})
    
    with tempfile.TemporaryDirectory() as tmpdir:
        chart_paths = generate_charts(analysis, tmpdir)
        
        pdf = PDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        
        pdf.set_font("helvetica", "B", 12)
        pdf.cell(0, 10, "Summary", ln=True)
        pdf.set_font("helvetica", "", 10)
        pdf.cell(0, 8, f"Total Videos Analyzed: {summary.get('total_videos', 0)}", ln=True)
        pdf.cell(0, 8, f"Total Comments Processed: {summary.get('total_comments_analyzed', 0)}", ln=True)
        pdf.ln(5)
        
        pdf.set_font("helvetica", "B", 12)
        pdf.cell(0, 10, "Methodology", ln=True)
        pdf.set_font("helvetica", "", 10)
        pdf.multi_cell(0, 8, "Data extracted from YouTube Data API dynamically based on user search query. Comments processed via NLP, cleaned, and analyzed for named entities using Spacy. Overall sentiment analyzed via VADER.")
        pdf.ln(5)

        pdf.set_font("helvetica", "B", 12)
        pdf.cell(0, 10, "Top topics/keywords", ln=True)
        pdf.set_font("helvetica", "", 10)
        for t in analysis.get("topics", []):
            pdf.cell(0, 8, f"{t.get('topic')}: {t.get('count')}", ln=True)
        if 'topics' in chart_paths:
            pdf.image(chart_paths['topics'], w=140)
        pdf.ln(5)

        pdf.set_font("helvetica", "B", 12)
        pdf.cell(0, 10, "Named entities frequency", ln=True)
        pdf.set_font("helvetica", "", 10)
        for k, v in analysis.get("organizations", {}).items():
            pdf.cell(0, 8, f"ORG: {k} ({v})", ln=True)
        if 'org' in chart_paths:
            pdf.image(chart_paths['org'], w=120)
        for k, v in analysis.get("persons", {}).items():
            pdf.cell(0, 8, f"PERSON: {k} ({v})", ln=True)
        if 'person' in chart_paths:
            pdf.image(chart_paths['person'], w=120)
        pdf.ln(5)

        pdf.set_font("helvetica", "B", 12)
        pdf.cell(0, 10, "Sentiment charts", ln=True)
        pdf.set_font("helvetica", "", 10)
        for k, v in analysis.get("sentiments", {}).items():
            pdf.cell(0, 8, f"{str(k).capitalize()}: {v}", ln=True)
        if 'sentiment' in chart_paths:
            pdf.image(chart_paths['sentiment'], w=120)
        pdf.ln(5)

        pdf.set_font("helvetica", "B", 12)
        pdf.cell(0, 10, "Key insights", ln=True)
        pdf.set_font("helvetica", "", 10)
        pdf.cell(0, 8, "1. Discourse is varied depending on dominant entities.", ln=True)
        pdf.cell(0, 8, "2. Sentiment splits highlight community alignment.", ln=True)
            
        try:
            pdf_bytes = pdf.output()
            file_stream = io.BytesIO(pdf_bytes)
            file_stream.seek(0)
        except Exception as e:
            raise HTTPException(status_code=500, detail=str(e))
        
    return StreamingResponse(
        file_stream, 
        media_type="application/pdf",
        headers={"Content-Disposition": "attachment; filename=analysis_report.pdf"}
    )
